"""Professional export pipeline for Scouting AI reports."""

from __future__ import annotations

import html
import json
import re
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import Any

import markdown

from src.config import SCOUTING_DIR, project_relative
from src.ingestion.utils import to_jsonable
from src.reports.pdf_report import render_pdf_report
from src.scouting.scouting_history import build_scouting_history_record, record_scouting_generation


def save_scouting_export(
    result: dict[str, Any],
    include_html: bool = False,
    include_docx: bool = False,
    include_pdf: bool = False,
    record_history: bool = True,
    use_api: bool | None = None,
) -> dict[str, Any]:
    """Save a professional scouting export and optionally record it in DuckDB."""

    SCOUTING_DIR.mkdir(parents=True, exist_ok=True)
    match_a = int(result["match_id_a"])
    player_a = int(result["player_id_a"])
    match_b = result.get("match_id_b")
    player_b = result.get("player_id_b")
    exported_at, suffix, paths = _build_paths(match_a, player_a, match_b, player_b)

    markdown_text = render_scouting_markdown(result)
    html_text = render_scouting_html(result, markdown_text)

    paths["markdown"].write_text(markdown_text, encoding="utf-8")
    with paths["json"].open("w", encoding="utf-8") as file:
        json.dump(to_jsonable(result), file, ensure_ascii=False, indent=2)
        file.write("\n")

    save_result: dict[str, Any] = {
        "markdown": _public_path(paths["markdown"]),
        "html": None,
        "json": _public_path(paths["json"]),
        "pdf": None,
        "docx": None,
        "exported_at": exported_at.isoformat(timespec="seconds"),
        "exported_at_utc": datetime.now(timezone.utc).replace(tzinfo=None).isoformat(timespec="seconds"),
        "export_suffix": suffix,
        "html_status": "not_requested",
        "pdf_status": "not_requested",
        "docx_status": "not_requested",
        "pdf_error_message": None,
        "pdf_warning_message": None,
        "docx_error_message": None,
        "pdf_result": {
            "status": "not_requested",
            "path": _public_path(paths["pdf"]),
            "error_message": None,
            "warning_message": None,
        },
        "docx_result": {
            "status": "not_requested",
            "path": _public_path(paths["docx"]),
            "error_message": None,
        },
    }

    if include_html:
        paths["html"].write_text(html_text, encoding="utf-8")
        save_result["html"] = _public_path(paths["html"])
        save_result["html_status"] = "generated"

    if include_docx:
        docx_result = render_scouting_docx(result, markdown_text, paths["docx"].as_posix())
        docx_result["path"] = _public_path(docx_result.get("path") or paths["docx"])
        save_result["docx_result"] = docx_result
        save_result["docx_status"] = docx_result["status"]
        save_result["docx_error_message"] = docx_result.get("error_message")
        if docx_result["status"] == "generated":
            save_result["docx"] = docx_result["path"]

    if include_pdf:
        pdf_result = render_pdf_report(html_text, paths["pdf"].as_posix())
        pdf_result["path"] = _public_path(pdf_result.get("path") or paths["pdf"])
        save_result["pdf_result"] = pdf_result
        save_result["pdf_status"] = pdf_result["status"]
        save_result["pdf_error_message"] = pdf_result.get("error_message")
        save_result["pdf_warning_message"] = pdf_result.get("warning_message")
        if pdf_result["status"] == "generated":
            save_result["pdf"] = pdf_result["path"]

    if record_history:
        selected_use_api = bool(result.get("use_api")) if use_api is None else bool(use_api)
        history_record = build_scouting_history_record(result, save_result, selected_use_api)
        save_result["generated_by"] = history_record.get("generated_by")
        try:
            record_scouting_generation(history_record)
            save_result["history_status"] = history_record.get("status")
            save_result["history_error_message"] = None
        except Exception as exc:
            save_result["history_status"] = "failed"
            save_result["history_error_message"] = str(exc)

    return save_result


def render_scouting_markdown(result: dict[str, Any]) -> str:
    summary = result.get("context_summary", {})
    warnings = result.get("warnings", [])
    language_warnings = result.get("language_warnings", [])
    narrative = _strip_top_heading(str(result.get("narrative_markdown") or ""))

    lines = [
        "# Reporte profesional de scouting",
        "",
        "## Datos generales",
        "",
        f"- **Modo:** {result.get('mode', 'individual')}",
        f"- **Jugador A:** {summary.get('player_a', 'N/D')} ({summary.get('team_a', 'N/D')})",
        f"- **Partido A:** {summary.get('match_a', 'N/D')}",
        f"- **Match ID A:** {result.get('match_id_a')}",
        f"- **Player ID A:** {result.get('player_id_a')}",
    ]
    if result.get("mode") == "comparativo":
        lines.extend(
            [
                f"- **Jugador B:** {summary.get('player_b', 'N/D')} ({summary.get('team_b', 'N/D')})",
                f"- **Partido B:** {summary.get('match_b', 'N/D')}",
                f"- **Match ID B:** {result.get('match_id_b')}",
                f"- **Player ID B:** {result.get('player_id_b')}",
            ]
        )

    lines.extend(
        [
            f"- **Estado narrativo:** {result.get('status')}",
            f"- **Modelo:** {result.get('model')}",
            f"- **Generado en:** {result.get('generated_at')}",
            "",
            narrative,
            "",
            "## Advertencias de lenguaje/factualidad",
            "",
        ]
    )

    all_warnings = [f"Factualidad/contexto: {warning}" for warning in warnings]
    all_warnings.extend(f"Lenguaje: {warning}" for warning in language_warnings)
    if all_warnings:
        lines.extend(f"- {warning}" for warning in all_warnings)
    else:
        lines.append("- No se detectaron advertencias de lenguaje ni factualidad.")

    lines.extend(
        [
            "",
            "## Trazabilidad",
            "",
            "- Fuente: StatsBomb Open Data.",
            "- Contexto derivado del comparador de jugadores y métricas visuales del proyecto.",
            "- El reporte evita inferencias no sustentadas sobre futuro, fichajes o valor de mercado.",
            "- Los datos raw permanecen sin modificar en `data/raw/`.",
            "- Historial persistente: `data/analytics/scouting_history.duckdb`.",
            "",
        ]
    )
    return "\n".join(lines)


def render_scouting_html(result: dict[str, Any], markdown_text: str | None = None) -> str:
    markdown_text = markdown_text if markdown_text is not None else render_scouting_markdown(result)
    body = markdown.markdown(markdown_text, extensions=["tables", "sane_lists"])
    title = html.escape(_html_title(result))
    return f"""<!doctype html>
<html lang="es">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{title}</title>
  <style>
    :root {{
      color-scheme: light;
      --ink: #17212b;
      --muted: #5f6b76;
      --line: #dce3ea;
      --accent: #0d6b5f;
      --accent-soft: #edf7f5;
    }}
    body {{
      margin: 0;
      background: #f5f7fa;
      color: var(--ink);
      font-family: "Segoe UI", Arial, sans-serif;
      line-height: 1.58;
    }}
    main {{
      max-width: 980px;
      min-height: 100vh;
      margin: 0 auto;
      padding: 42px 28px 68px;
      background: #fff;
      box-shadow: 0 0 0 1px rgba(23, 33, 43, 0.07);
    }}
    h1 {{
      margin: 0 0 26px;
      padding-bottom: 14px;
      border-bottom: 4px solid var(--accent);
      font-size: 32px;
    }}
    h2 {{
      margin-top: 32px;
      padding-bottom: 8px;
      border-bottom: 1px solid var(--line);
      color: #123c38;
      font-size: 22px;
    }}
    h3 {{ margin-top: 22px; color: #243447; }}
    p, li {{ color: var(--ink); }}
    strong {{ color: #0d171f; }}
    ul {{ padding-left: 22px; }}
    code {{
      background: #eef2f6;
      border-radius: 4px;
      padding: 2px 5px;
    }}
    table {{
      width: 100%;
      border-collapse: collapse;
      margin: 16px 0 24px;
      font-size: 14px;
    }}
    th, td {{
      border: 1px solid var(--line);
      padding: 8px 10px;
      text-align: left;
      vertical-align: top;
    }}
    th {{
      background: var(--accent-soft);
      color: #0c3f38;
    }}
  </style>
</head>
<body>
  <main>
    {body}
  </main>
</body>
</html>
"""


def render_scouting_docx(result: dict[str, Any], markdown_text: str, output_path: str) -> dict[str, Any]:
    path = Path(output_path)
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt, RGBColor

        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

        styles = document.styles
        styles["Normal"].font.name = "Arial"
        styles["Normal"].font.size = Pt(10)
        for style_name in ("Heading 1", "Heading 2", "Heading 3"):
            style = styles[style_name]
            style.font.name = "Arial"
            style.font.color.rgb = RGBColor(13, 107, 95)

        _add_markdown_to_docx(document, markdown_text)
        title = document.paragraphs[0] if document.paragraphs else None
        if title is not None:
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        path.parent.mkdir(parents=True, exist_ok=True)
        document.save(path)
        return {"status": "generated", "path": path.as_posix(), "error_message": None}
    except Exception as exc:
        return {"status": "failed", "path": path.as_posix(), "error_message": str(exc)}


def _build_paths(
    match_a: int,
    player_a: int,
    match_b: Any,
    player_b: Any,
) -> tuple[datetime, str, dict[str, Path]]:
    exported_at = datetime.now()
    while True:
        suffix = exported_at.strftime("%Y%m%d_%H%M%S")
        if match_b is not None and player_b is not None:
            base_name = f"scouting.match-{match_a}.{player_a}_vs_match-{int(match_b)}.{int(player_b)}_{suffix}"
        else:
            base_name = f"scouting.match-{match_a}.{player_a}_{suffix}"
        paths = {
            "markdown": SCOUTING_DIR / f"{base_name}.md",
            "html": SCOUTING_DIR / f"{base_name}.html",
            "json": SCOUTING_DIR / f"{base_name}.json",
            "pdf": SCOUTING_DIR / f"{base_name}.pdf",
            "docx": SCOUTING_DIR / f"{base_name}.docx",
        }
        if not any(path.exists() for path in paths.values()):
            return exported_at, suffix, paths
        exported_at += timedelta(seconds=1)


def _add_markdown_to_docx(document: Any, markdown_text: str) -> None:
    for line in markdown_text.splitlines():
        clean = line.strip()
        if not clean:
            continue
        if clean.startswith("# "):
            document.add_heading(_strip_markdown(clean[2:]), level=0)
        elif clean.startswith("## "):
            document.add_heading(_strip_markdown(clean[3:]), level=1)
        elif clean.startswith("### "):
            document.add_heading(_strip_markdown(clean[4:]), level=2)
        elif clean.startswith("- "):
            document.add_paragraph(_strip_markdown(clean[2:]), style="List Bullet")
        else:
            document.add_paragraph(_strip_markdown(clean))


def _strip_top_heading(markdown_text: str) -> str:
    lines = markdown_text.strip().splitlines()
    if lines and lines[0].startswith("# "):
        return "\n".join(lines[1:]).strip()
    return markdown_text.strip()


def _strip_markdown(value: str) -> str:
    clean = re.sub(r"^#+\s*", "", value)
    clean = clean.replace("**", "")
    clean = clean.replace("`", "")
    return clean


def _html_title(result: dict[str, Any]) -> str:
    summary = result.get("context_summary", {})
    player_a = summary.get("player_a") or result.get("player_id_a")
    if result.get("mode") == "comparativo":
        player_b = summary.get("player_b") or result.get("player_id_b")
        return f"Scouting {player_a} vs {player_b}"
    return f"Scouting {player_a}"


def _public_path(path_value: str | Path) -> str:
    path = path_value if isinstance(path_value, Path) else Path(str(path_value))
    if not path.is_absolute():
        return path.as_posix()
    return project_relative(path)
