"""DOCX rendering for final match reports."""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from src.reports.branding import add_docx_footer


def render_docx_report(report: dict[str, Any], output_path: str) -> dict[str, Any]:
    path = Path(output_path)
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt, RGBColor

        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

        styles = document.styles
        styles["Normal"].font.name = "Arial"
        styles["Normal"].font.size = Pt(10)
        for style_name in ("Heading 1", "Heading 2", "Heading 3"):
            style = styles[style_name]
            style.font.name = "Arial"
            style.font.color.rgb = RGBColor(15, 90, 84)

        title = document.add_heading("Reporte del partido", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        _add_metadata(document, report)
        _add_executive_summary(document, report)
        _add_team_stats(document, report)
        _add_dominance(document, report)
        _add_key_moments(document, report)
        _add_impact_players(document, report)
        _add_narrative(document, report)
        _add_quality(document, report)
        _add_validation(document, report)
        _add_traceability(document, report)
        add_docx_footer(document)

        path.parent.mkdir(parents=True, exist_ok=True)
        document.save(str(path))
        return {
            "status": "generated",
            "path": path.as_posix(),
            "error_message": None,
        }
    except Exception as exc:
        return {
            "status": "failed",
            "path": path.as_posix(),
            "error_message": str(exc),
        }


def _add_metadata(document: Any, report: dict[str, Any]) -> None:
    summary = report.get("match_summary", {})
    metadata = report.get("match_metadata", {})
    document.add_heading("Datos generales", level=1)
    rows = [
        ("Competencia", metadata.get("competition_name")),
        ("Temporada", metadata.get("season_name")),
        ("Fecha", summary.get("match_date") or metadata.get("match_date")),
        ("Partido", _score_line(summary)),
        ("Marcador", f"{summary.get('home_score')}-{summary.get('away_score')}"),
        ("Match ID", summary.get("match_id") or report.get("match_id")),
        ("Estadio", metadata.get("stadium_name")),
        ("Arbitro", metadata.get("referee_name")),
    ]
    _add_key_value_table(document, rows)


def _add_executive_summary(document: Any, report: dict[str, Any]) -> None:
    summary = report.get("match_summary", {})
    analytics = report.get("analytics", {})
    dominance = analytics.get("dominance", [])
    leader = dominance[0] if dominance else {}
    document.add_heading("Resumen ejecutivo", level=1)
    document.add_paragraph(
        f"{_score_line(summary)}. El ganador fue {summary.get('winner_team_name', 'N/D')}. "
        f"El dominio estimado favoreció a {leader.get('team_name', 'N/D')} "
        f"con score {leader.get('dominance_score', 'N/D')}. El reporte consolida "
        "estadística, análisis avanzado, narrativa AI, calidad narrativa y trazabilidad."
    )


def _add_team_stats(document: Any, report: dict[str, Any]) -> None:
    analytics = report.get("analytics", {})
    team_stats = analytics.get("team_stats", [])
    dangerous_counts = _dangerous_counts(analytics.get("dangerous_attacks", []))
    possession_summary = analytics.get("possession_summary", {})
    possessions_by_team = (
        possession_summary.get("possessions_by_team", {}) if isinstance(possession_summary, dict) else {}
    )
    document.add_heading("Estadísticas principales", level=1)
    rows = []
    for team in team_stats:
        name = str(team.get("team_name"))
        rows.append(
            [
                name,
                team.get("shots", 0),
                team.get("goals", 0),
                _number(team.get("xg")),
                team.get("passes", 0),
                f"{_number(team.get('pass_completion_pct'))}%",
                possessions_by_team.get(name, "N/D"),
                dangerous_counts.get(name, 0),
            ]
        )
    _add_table(
        document,
        ["Equipo", "Tiros", "Goles", "xG", "Pases", "Precisión", "Posesiones", "Ataques peligrosos"],
        rows,
    )


def _add_dominance(document: Any, report: dict[str, Any]) -> None:
    analytics = report.get("analytics", {})
    summary = report.get("match_summary", {})
    dominance = analytics.get("dominance", [])
    xg_breakdown = analytics.get("xg_breakdown", [])
    document.add_heading("Lectura del dominio", level=1)
    if not dominance:
        document.add_paragraph("No hay datos de dominio disponibles.")
        return
    leader = dominance[0]
    xg_text = ", ".join(f"{row.get('team_name')} xG {row.get('xg_total')}" for row in xg_breakdown)
    document.add_paragraph(
        f"El equipo dominante por volumen fue {leader.get('team_name')}, con "
        f"{leader.get('shots')} tiros, {leader.get('final_third_entries')} entradas al último tercio "
        f"y score de dominio {leader.get('dominance_score')}. En xG: {xg_text}. "
        f"El ganador fue {summary.get('winner_team_name')}, señal de efectividad frente al dominio territorial."
    )


def _add_key_moments(document: Any, report: dict[str, Any]) -> None:
    document.add_heading("Momentos clave", level=1)
    key_moments = report.get("analytics", {}).get("key_moments", [])
    if not key_moments:
        document.add_paragraph("No se detectaron momentos clave.")
        return
    rows = []
    for moment in key_moments:
        second = int(moment.get("second") or 0)
        rows.append([f"{moment.get('minute')}:{second:02d}", moment.get("type"), moment.get("title")])
    _add_table(document, ["Minuto", "Tipo", "Descripción"], rows)


def _add_impact_players(document: Any, report: dict[str, Any]) -> None:
    document.add_heading("Jugadores destacados", level=1)
    players = report.get("analytics", {}).get("impact_players", [])
    if not players:
        document.add_paragraph("No hay jugadores destacados disponibles.")
        return
    rows = []
    for player in players[:10]:
        rows.append(
            [
                player.get("player_name"),
                player.get("team_name"),
                player.get("impact_score"),
                player.get("goals"),
                player.get("assists"),
                player.get("shots"),
                player.get("xg"),
                player.get("key_passes"),
            ]
        )
    _add_table(document, ["Jugador", "Equipo", "Score", "Goles", "Asist.", "Tiros", "xG", "Pases clave"], rows)


def _add_narrative(document: Any, report: dict[str, Any]) -> None:
    document.add_heading("Narración AI", level=1)
    narrative = str(report.get("narrative", {}).get("narrative_markdown") or "Narración no disponible.")
    for line in narrative.splitlines():
        clean = line.strip()
        if not clean:
            continue
        if clean.startswith("# "):
            document.add_heading(_strip_markdown(clean), level=2)
        elif clean.startswith("## "):
            document.add_heading(_strip_markdown(clean), level=3)
        elif clean.startswith("- "):
            document.add_paragraph(_strip_markdown(clean[2:]), style="List Bullet")
        else:
            document.add_paragraph(_strip_markdown(clean))


def _add_quality(document: Any, report: dict[str, Any]) -> None:
    document.add_heading("Evaluación de calidad narrativa", level=1)
    quality = report.get("quality", {})
    rows = [
        ["Overall", quality.get("overall_score")],
        ["Factualidad", quality.get("factuality_score")],
        ["Cobertura", quality.get("coverage_score")],
        ["Claridad", quality.get("clarity_score")],
        ["Emoción", quality.get("excitement_score")],
        ["Profundidad táctica", quality.get("tactical_depth_score")],
    ]
    _add_table(document, ["Métrica", "Score"], rows)
    warnings = quality.get("warnings", [])
    if warnings:
        document.add_heading("Warnings", level=2)
        for warning in warnings:
            document.add_paragraph(str(warning), style="List Bullet")


def _add_validation(document: Any, report: dict[str, Any]) -> None:
    document.add_heading("Validación futbolística", level=1)
    validation = report.get("analytics", {}).get("validation", {})
    document.add_paragraph(f"Status: {validation.get('status', 'N/D')}")
    findings = validation.get("findings", [])
    if findings:
        rows = [[f.get("severity"), f.get("message"), f.get("rows")] for f in findings]
        _add_table(document, ["Severidad", "Hallazgo", "Filas"], rows)
    else:
        document.add_paragraph("Sin hallazgos.")


def _add_traceability(document: Any, report: dict[str, Any]) -> None:
    document.add_heading("Trazabilidad", level=1)
    items = [
        "Fuente: StatsBomb Open Data.",
        "Raw JSON preservado en data/raw/ sin modificaciones analíticas.",
        "DuckDB analítico: data/analytics/statsbomb.duckdb.",
        "Contexto AI curado desde src/analytics/ai_context.py.",
        f"Fecha de generación: {report.get('generated_at')}.",
    ]
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def _add_key_value_table(document: Any, rows: list[tuple[str, Any]]) -> None:
    _add_table(document, ["Campo", "Valor"], [[key, _value(value)] for key, value in rows])


def _add_table(document: Any, headers: list[str], rows: list[list[Any]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        paragraph = header_cells[index].paragraphs[0]
        paragraph.text = str(header)
        for run in paragraph.runs:
            run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = _value(value)
    document.add_paragraph()


def _score_line(summary: dict[str, Any]) -> str:
    return (
        f"{summary.get('home_team_name')} {summary.get('home_score')}-"
        f"{summary.get('away_score')} {summary.get('away_team_name')}"
    )


def _dangerous_counts(attacks: list[dict[str, Any]]) -> dict[str, int]:
    counts: dict[str, int] = {}
    for attack in attacks:
        team = str(attack.get("team_name") or "Sin equipo")
        counts[team] = counts.get(team, 0) + 1
    return counts


def _number(value: Any) -> str:
    if value is None:
        return "N/D"
    try:
        return f"{float(value):.2f}"
    except (TypeError, ValueError):
        return str(value)


def _value(value: Any) -> str:
    return str(value) if value not in (None, "") else "N/D"


def _strip_markdown(value: str) -> str:
    clean = re.sub(r"^#+\s*", "", value)
    clean = clean.replace("**", "")
    clean = clean.replace("`", "")
    return clean
