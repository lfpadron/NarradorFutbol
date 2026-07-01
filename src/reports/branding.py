"""Shared report branding for generated PDF and DOCX artifacts."""

from __future__ import annotations

import base64
from functools import lru_cache
from pathlib import Path
from typing import Any

from src.config import PROJECT_ROOT

FOOTER_TEXT = "Construido por Luis Fernando Padrón"
FOOTER_LOGO_PATH = PROJECT_ROOT / "app" / "assets" / "astrogato_footer_logo.png"


@lru_cache(maxsize=1)
def footer_logo_data_uri() -> str:
    if not FOOTER_LOGO_PATH.exists():
        return ""
    encoded = base64.b64encode(FOOTER_LOGO_PATH.read_bytes()).decode("ascii")
    return f"data:image/png;base64,{encoded}"


def footer_logo_path() -> Path | None:
    return FOOTER_LOGO_PATH if FOOTER_LOGO_PATH.exists() else None


def add_docx_footer(document: Any) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    logo_path = footer_logo_path()
    for section in document.sections:
        section.footer_distance = Inches(0.2)
        paragraph = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.text = ""
        if logo_path is not None:
            logo_run = paragraph.add_run()
            logo_run.add_picture(str(logo_path), width=Inches(0.28))
            paragraph.add_run("  ")
        text_run = paragraph.add_run(FOOTER_TEXT)
        text_run.font.name = "Arial"
        text_run.font.size = Pt(8)
        text_run.font.bold = True
        text_run.font.color.rgb = RGBColor(71, 85, 105)


def add_report_footer_to_html(html: str) -> str:
    logo_uri = footer_logo_data_uri()
    logo_markup = (
        f'<img src="{logo_uri}" alt="Astrogato" class="astrogato-report-footer__logo">'
        if logo_uri
        else ""
    )
    style = """
<style>
  @page {
    margin-bottom: 76px;
    @bottom-center {
      content: element(astrogato-report-footer);
    }
  }
  .astrogato-report-footer {
    position: running(astrogato-report-footer);
    color: #475569;
    font-family: "Segoe UI", Arial, sans-serif;
    font-size: 9px;
    font-weight: 700;
    text-align: center;
    white-space: nowrap;
  }
  .astrogato-report-footer__logo {
    width: 28px;
    height: 28px;
    border-radius: 50%;
    object-fit: cover;
    vertical-align: middle;
    margin-right: 8px;
  }
  @media screen {
    .astrogato-report-footer {
      margin-top: 28px;
      padding-top: 12px;
      border-top: 1px solid #d9e0e7;
    }
  }
</style>
"""
    footer = f"""
<footer class="astrogato-report-footer">
  {logo_markup}<span>{FOOTER_TEXT}</span>
</footer>
"""
    if "</head>" in html:
        html = html.replace("</head>", f"{style}\n</head>", 1)
    else:
        html = f"{style}\n{html}"
    if "</body>" in html:
        return html.replace("</body>", f"{footer}\n</body>", 1)
    return f"{html}\n{footer}"


def draw_reportlab_footer(canvas: Any, document: Any) -> None:
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    from reportlab.lib.utils import ImageReader

    page_width, _ = document.pagesize
    font_name = "Helvetica-Bold"
    font_size = 8.5
    logo_size = 0.28 * inch
    gap = 0.08 * inch
    y = 0.34 * inch
    text_width = canvas.stringWidth(FOOTER_TEXT, font_name, font_size)
    total_width = text_width + gap + (logo_size if footer_logo_path() else 0)
    x = (page_width - total_width) / 2

    canvas.saveState()
    canvas.setFillColor(colors.HexColor("#475569"))
    canvas.setFont(font_name, font_size)
    logo_path = footer_logo_path()
    if logo_path is not None:
        canvas.drawImage(
            ImageReader(str(logo_path)),
            x,
            y - logo_size / 2,
            width=logo_size,
            height=logo_size,
            mask="auto",
        )
        x += logo_size + gap
    canvas.drawString(x, y - (font_size / 2.8), FOOTER_TEXT)
    canvas.restoreState()
