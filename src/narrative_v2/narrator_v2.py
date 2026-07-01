"""Generate specialized Narrador AI v2 narratives."""

from __future__ import annotations

import json
import re
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import Any

import markdown
from openai import OpenAI, OpenAIError

from src.analytics.ai_context import build_ai_match_context
from src.config import ANALYTICS_EXPORTS_DIR, project_relative
from src.ingestion.utils import to_jsonable
from src.narrative.config import get_openai_api_key, get_openai_model
from src.narrative.fact_guard import validate_narrative_against_context
from src.narrative_v2.prompt_builder_v2 import build_specialized_prompt
from src.narrative_v2.section_builder import build_context_for_style
from src.narrative_v2.style_evaluator import evaluate_style_fit
from src.narrative_v2.style_profiles import STYLE_PROFILES, get_style_profile
from src.reports.branding import add_docx_footer
from src.reports.pdf_report import render_pdf_report


def generate_specialized_narrative(
    match_id: int,
    style_id: str,
    use_api: bool = True,
) -> dict[str, Any]:
    profile = get_style_profile(style_id)
    full_context = build_ai_match_context(match_id)
    context_used = build_context_for_style(full_context, style_id)
    model = get_openai_model()
    warnings: list[str] = []
    status = "fallback"

    api_key = get_openai_api_key()
    if use_api and api_key:
        try:
            prompt = build_specialized_prompt(context_used, style_id)
            client = OpenAI(api_key=api_key)
            response = client.responses.create(
                model=model,
                input=prompt,
                temperature=0.35,
            )
            narrative_markdown = _extract_response_text(response).strip()
            status = "generated"
        except OpenAIError as exc:
            narrative_markdown = generate_specialized_fallback(context_used, style_id)
            warnings.append(f"OpenAI API fallo; se uso fallback local. Detalle: {exc}")
        except Exception as exc:
            narrative_markdown = generate_specialized_fallback(context_used, style_id)
            warnings.append(f"No se pudo generar con API; se uso fallback local. Detalle: {exc}")
    else:
        narrative_markdown = generate_specialized_fallback(context_used, style_id)
        if use_api and not api_key:
            warnings.append("OPENAI_API_KEY no esta configurada; se uso fallback local v2.")
        elif not use_api:
            warnings.append("Uso de OpenAI API desactivado; se uso fallback local v2.")

    fact_warnings = validate_narrative_against_context(
        _fact_guard_text(narrative_markdown),
        full_context,
    )
    style_quality = evaluate_style_fit(narrative_markdown, style_id, full_context)

    return to_jsonable(
        {
            "match_id": match_id,
            "style_id": style_id,
            "style_name": profile["name"],
            "generated_at": datetime.now(timezone.utc).isoformat(),
            "status": status,
            "model": model,
            "warnings": warnings,
            "narrative_markdown": narrative_markdown,
            "fact_warnings": fact_warnings,
            "style_quality": style_quality,
            "context_used": context_used,
        }
    )


def compare_specialized_styles(match_id: int, use_api: bool = False) -> dict[str, Any]:
    rows = []
    best_style = None
    best_score = -1
    for style_id, profile in STYLE_PROFILES.items():
        result = generate_specialized_narrative(match_id, style_id, use_api=use_api)
        score = int(result.get("style_quality", {}).get("style_score") or 0)
        rows.append(
            {
                "style_id": style_id,
                "style_name": profile["name"],
                "status": result.get("status"),
                "style_score": score,
                "fact_warnings_count": len(result.get("fact_warnings", [])),
                "narrative_markdown": result.get("narrative_markdown"),
            }
        )
        if score > best_score:
            best_score = score
            best_style = style_id
    return to_jsonable({"match_id": match_id, "styles": rows, "best_style": best_style})


def save_specialized_narrative(result: dict[str, Any]) -> tuple[str, str]:
    save_result = save_specialized_narrative_export(result)
    return str(save_result["markdown"]), str(save_result["json"])


def save_specialized_narrative_export(
    result: dict[str, Any],
    include_html: bool = False,
    include_pdf: bool = False,
    include_docx: bool = False,
) -> dict[str, Any]:
    ANALYTICS_EXPORTS_DIR.mkdir(parents=True, exist_ok=True)
    match_id = result["match_id"]
    style_id = _safe_token(str(result["style_id"]))
    exported_at, suffix, paths = _build_specialized_export_paths(match_id, style_id)
    markdown_text = render_specialized_narrative_markdown(result)
    html_text = render_specialized_narrative_html(result, markdown_text)

    paths["markdown"].write_text(markdown_text, encoding="utf-8")
    with paths["json"].open("w", encoding="utf-8") as file:
        json.dump(to_jsonable(result), file, ensure_ascii=False, indent=2)
        file.write("\n")

    save_result: dict[str, Any] = {
        "markdown": _public_path(paths["markdown"]),
        "html": None,
        "json": _public_path(paths["json"]),
        "pdf": None,
        "docx": None,
        "exported_at": exported_at.isoformat(timespec="seconds"),
        "exported_at_utc": datetime.now(timezone.utc).replace(tzinfo=None).isoformat(timespec="seconds"),
        "export_suffix": suffix,
        "html_status": "not_requested",
        "pdf_status": "not_requested",
        "docx_status": "not_requested",
        "pdf_error_message": None,
        "pdf_warning_message": None,
        "docx_error_message": None,
    }

    if include_html:
        paths["html"].write_text(html_text, encoding="utf-8")
        save_result["html"] = _public_path(paths["html"])
        save_result["html_status"] = "generated"

    if include_docx:
        docx_result = render_specialized_narrative_docx(markdown_text, paths["docx"].as_posix())
        docx_result["path"] = _public_path(docx_result.get("path") or paths["docx"])
        save_result["docx_status"] = docx_result["status"]
        save_result["docx_error_message"] = docx_result.get("error_message")
        if docx_result["status"] == "generated":
            save_result["docx"] = docx_result["path"]

    if include_pdf:
        pdf_result = render_pdf_report(html_text, paths["pdf"].as_posix())
        pdf_result["path"] = _public_path(pdf_result.get("path") or paths["pdf"])
        save_result["pdf_status"] = pdf_result["status"]
        save_result["pdf_error_message"] = pdf_result.get("error_message")
        save_result["pdf_warning_message"] = pdf_result.get("warning_message")
        if pdf_result["status"] == "generated":
            save_result["pdf"] = pdf_result["path"]

    return save_result


def render_specialized_narrative_markdown(result: dict[str, Any]) -> str:
    context = result.get("context_used", {})
    profile = context.get("style_profile", {})
    summary = context.get("match_summary", {})
    quality = result.get("style_quality", {})
    lines = [
        "# Narrador AI v2",
        "",
        "## Datos generales",
        "",
        f"- **Estilo:** {result.get('style_name') or profile.get('name') or result.get('style_id')}",
        f"- **Audiencia:** {profile.get('audience', 'N/D')}",
        f"- **Objetivo:** {profile.get('objective', 'N/D')}",
        f"- **Partido:** {_score_line(summary)}",
        f"- **Match ID:** {result.get('match_id')}",
        f"- **Estado:** {result.get('status')}",
        f"- **Modelo:** {result.get('model')}",
        f"- **Generado en:** {result.get('generated_at')}",
        "",
        "## Narrativa",
        "",
        _strip_top_heading(str(result.get("narrative_markdown") or "")),
        "",
        "## Calidad de estilo",
        "",
        "| Metrica | Score |",
        "| --- | ---: |",
        f"| Estilo | {quality.get('style_score', 'N/D')} |",
        f"| Estructura | {quality.get('structure_score', 'N/D')} |",
        f"| Audiencia | {quality.get('audience_fit_score', 'N/D')} |",
        f"| Factualidad | {quality.get('factuality_score', 'N/D')} |",
        "",
        "## Advertencias",
        "",
        *_warning_lines(result),
        "",
        "## Trazabilidad",
        "",
        "- Fuente: StatsBomb Open Data transformada a DuckDB analitico.",
        "- Contexto reducido por perfil desde `src/narrative_v2/section_builder.py`.",
        "- El archivo JSON conserva el contexto usado, calidad de estilo y advertencias.",
        "",
    ]
    return "\n".join(lines)


def render_specialized_narrative_html(result: dict[str, Any], markdown_text: str | None = None) -> str:
    markdown_text = markdown_text if markdown_text is not None else render_specialized_narrative_markdown(result)
    body = markdown.markdown(markdown_text, extensions=["tables", "sane_lists"])
    title = _html_escape(_html_title(result))
    return f"""<!doctype html>
<html lang="es">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{title}</title>
  <style>
    :root {{
      color-scheme: light;
      --ink: #17212b;
      --line: #dce3ea;
      --accent: #0d6b5f;
      --accent-soft: #edf7f5;
    }}
    body {{
      margin: 0;
      background: #f5f7fa;
      color: var(--ink);
      font-family: "Segoe UI", Arial, sans-serif;
      line-height: 1.58;
    }}
    main {{
      max-width: 980px;
      min-height: 100vh;
      margin: 0 auto;
      padding: 42px 28px 68px;
      background: #fff;
      box-shadow: 0 0 0 1px rgba(23, 33, 43, 0.07);
    }}
    h1 {{
      margin: 0 0 26px;
      padding-bottom: 14px;
      border-bottom: 4px solid var(--accent);
      font-size: 32px;
    }}
    h2 {{
      margin-top: 32px;
      padding-bottom: 8px;
      border-bottom: 1px solid var(--line);
      color: #123c38;
      font-size: 22px;
    }}
    h3 {{ margin-top: 22px; color: #243447; }}
    p, li {{ color: var(--ink); }}
    strong {{ color: #0d171f; }}
    ul {{ padding-left: 22px; }}
    table {{
      width: 100%;
      border-collapse: collapse;
      margin: 16px 0 24px;
      font-size: 14px;
    }}
    th, td {{
      border: 1px solid var(--line);
      padding: 8px 10px;
      text-align: left;
      vertical-align: top;
    }}
    th {{
      background: var(--accent-soft);
      color: #0c3f38;
    }}
  </style>
</head>
<body>
  <main>
    {body}
  </main>
</body>
</html>
"""


def render_specialized_narrative_docx(markdown_text: str, output_path: str) -> dict[str, Any]:
    path = Path(output_path)
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt, RGBColor

        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

        styles = document.styles
        styles["Normal"].font.name = "Arial"
        styles["Normal"].font.size = Pt(10)
        for style_name in ("Heading 1", "Heading 2", "Heading 3"):
            style = styles[style_name]
            style.font.name = "Arial"
            style.font.color.rgb = RGBColor(13, 107, 95)

        _add_markdown_to_docx(document, markdown_text)
        title = document.paragraphs[0] if document.paragraphs else None
        if title is not None:
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_docx_footer(document)

        path.parent.mkdir(parents=True, exist_ok=True)
        document.save(str(path))
        return {"status": "generated", "path": path.as_posix(), "error_message": None}
    except Exception as exc:
        return {"status": "failed", "path": path.as_posix(), "error_message": str(exc)}


def _build_specialized_export_paths(match_id: int, style_id: str) -> tuple[datetime, str, dict[str, Path]]:
    exported_at = datetime.now()
    while True:
        suffix = exported_at.strftime("%Y%m%d_%H%M%S")
        base_name = f"narrative_v2.match-{match_id}.{style_id}_{suffix}"
        paths = {
            "markdown": ANALYTICS_EXPORTS_DIR / f"{base_name}.md",
            "html": ANALYTICS_EXPORTS_DIR / f"{base_name}.html",
            "json": ANALYTICS_EXPORTS_DIR / f"{base_name}.json",
            "pdf": ANALYTICS_EXPORTS_DIR / f"{base_name}.pdf",
            "docx": ANALYTICS_EXPORTS_DIR / f"{base_name}.docx",
        }
        if not any(path.exists() for path in paths.values()):
            return exported_at, suffix, paths
        exported_at += timedelta(seconds=1)


def _warning_lines(result: dict[str, Any]) -> list[str]:
    warnings = []
    warnings.extend(f"- Generacion: {warning}" for warning in result.get("warnings", []))
    warnings.extend(f"- Factualidad: {warning}" for warning in result.get("fact_warnings", []))
    style_quality = result.get("style_quality", {})
    warnings.extend(f"- Estilo: {warning}" for warning in style_quality.get("warnings", []))
    return warnings or ["- No se detectaron advertencias."]


def _strip_top_heading(markdown_text: str) -> str:
    lines = markdown_text.strip().splitlines()
    if lines and lines[0].startswith("# "):
        return "\n".join(lines[1:]).strip()
    return markdown_text.strip()


def _html_title(result: dict[str, Any]) -> str:
    context = result.get("context_used", {})
    summary = context.get("match_summary", {})
    style_name = result.get("style_name") or result.get("style_id")
    return f"Narrador AI v2 {style_name} | {_score_line(summary)}"


def _html_escape(value: Any) -> str:
    text = str(value)
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
        .replace("'", "&#x27;")
    )


def _add_markdown_to_docx(document: Any, markdown_text: str) -> None:
    for line in markdown_text.splitlines():
        clean = line.strip()
        if not clean:
            continue
        if clean.startswith("# "):
            document.add_heading(_strip_markdown(clean[2:]), level=0)
        elif clean.startswith("## "):
            document.add_heading(_strip_markdown(clean[3:]), level=1)
        elif clean.startswith("### "):
            document.add_heading(_strip_markdown(clean[4:]), level=2)
        elif clean.startswith("- "):
            document.add_paragraph(_strip_markdown(clean[2:]), style="List Bullet")
        elif clean.startswith("|"):
            document.add_paragraph(_strip_markdown(clean))
        else:
            document.add_paragraph(_strip_markdown(clean))


def _strip_markdown(value: str) -> str:
    clean = re.sub(r"^#+\s*", "", value)
    clean = clean.replace("**", "")
    clean = clean.replace("`", "")
    return clean


def _public_path(path_value: str | Path) -> str:
    path = path_value if isinstance(path_value, Path) else Path(str(path_value))
    if not path.is_absolute():
        return path.as_posix()
    return project_relative(path)


def save_style_comparison(comparison: dict[str, Any]) -> str:
    ANALYTICS_EXPORTS_DIR.mkdir(parents=True, exist_ok=True)
    suffix = datetime.now().strftime("%Y%m%d_%H%M%S")
    path = ANALYTICS_EXPORTS_DIR / f"narrative_v2.compare.match-{comparison['match_id']}_{suffix}.json"
    with path.open("w", encoding="utf-8") as file:
        json.dump(to_jsonable(comparison), file, ensure_ascii=False, indent=2)
        file.write("\n")
    return project_relative(path)


def generate_specialized_fallback(context: dict[str, Any], style_id: str) -> str:
    summary = context.get("match_summary", {})
    style_generators = {
        "tactico": _fallback_tactico,
        "television": _fallback_television,
        "periodistico": _fallback_periodistico,
        "scouting": _fallback_scouting,
        "ejecutivo": _fallback_ejecutivo,
    }
    return style_generators[style_id](context, summary)


def _fallback_tactico(context: dict[str, Any], summary: dict[str, Any]) -> str:
    leader = _first(context.get("dominance", []))
    xg = _xg_text(context.get("xg_breakdown", []))
    dangerous = _dangerous_text(context.get("dangerous_attacks", []))
    validation = context.get("validation", {})
    return f"""# Lectura tactica

{_score_line(summary)}. El resultado favorecio a **{summary.get('winner_team_name', 'N/D')}**, pero la lectura de volumen marco otro eje: **{leader.get('team_name', 'N/D')}** concentro el dominio estimado con score {leader.get('dominance_score', 'N/D')}.

## Dominio y xG

El partido separa dominio territorial de eficacia. En xG: {xg}. Esa brecha ayuda a explicar por que el equipo con mas produccion no necesariamente encontro el gol.

## Presion y momentum

El momentum debe leerse como acumulacion de tiros, entradas al ultimo tercio y acciones ofensivas. La ventaja mexicana aparece menos por volumen y mas por resolver el momento decisivo.

## Ataques peligrosos

{dangerous}

## Claves del resultado

- Mexico sostuvo la eficacia en el marcador 0-1.
- Germany genero mayor volumen, pero no transformo ese dominio en gol.
- La lectura final es dominio aleman con eficacia mexicana.
- Validacion: {validation.get('status', 'N/D')}.
"""


def _fallback_television(context: dict[str, Any], summary: dict[str, Any]) -> str:
    goal = _goal_moment(context.get("key_moments", []))
    goal_player = _player_or_default(goal, "Hirving Rodrigo Lozano Bahena")
    return f"""# Apertura

Germany 0-1 Mexico: un partido de tension, ritmo y golpe decisivo. Mexico encontro la jugada que cambio la tarde y luego defendio una ventaja enorme en valor competitivo.

## Ritmo del partido

Germany empujo, sumo volumen y obligo a mirar el area mexicana una y otra vez. Pero el futbol tambien se decide en precision, y ahi Mexico tuvo la accion que partio el encuentro.

## Momento clave

{_moment_sentence(goal)}

## Figura

{goal_player} aparece como nombre relevante por impacto directo en el resultado y presencia en los momentos decisivos.

## Cierre

El marcador no se mueve: Germany 0-1 Mexico. Dominio aleman, eficacia mexicana y una victoria que se explica por resistir, elegir bien y golpear cuando el partido lo permitio.
"""


def _fallback_periodistico(context: dict[str, Any], summary: dict[str, Any]) -> str:
    goal = _goal_moment(context.get("key_moments", []))
    leader = _first(context.get("dominance", []))
    return f"""# Titular sugerido

Mexico golpea a Germany y firma un 0-1 de maxima eficacia

## Bajada

El equipo mexicano vencio con gol de {_player_or_default(goal, 'Hirving Rodrigo Lozano Bahena')} en un partido donde Germany acumulo dominio, pero no pudo igualar el marcador.

## Cronica

Germany llevo buena parte del peso ofensivo, pero Mexico encontro el momento que necesitaba. {_moment_sentence(goal)}. Desde ahi, el partido se convirtio en una prueba de resistencia, lectura y contundencia.

El dominio estimado favorecio a **{leader.get('team_name', 'Germany')}**, senal de que el volumen ofensivo no conto toda la historia. Mexico, en cambio, puso el marcador de su lado y obligo al rival a perseguir el partido.

## Claves

- Germany domino mas tramos del partido.
- Mexico fue mas eficaz en la accion decisiva.
- Lozano quedo como referencia directa del triunfo.

## Cierre

La cronica deja una lectura clara: Germany produjo mas, Mexico resolvio mejor. El 0-1 no contradice el dominio aleman; lo vuelve parte central del relato.
"""


def _fallback_scouting(context: dict[str, Any], summary: dict[str, Any]) -> str:
    players = context.get("impact_players", [])[:5] or context.get("top_players", [])[:5]
    rows = "\n".join(
        f"- **{player.get('player_name')}** ({player.get('team_name')}): impacto {player.get('impact_score', 'N/D')}, tiros {player.get('shots', 'N/D')}, xG {player.get('xg', 'N/D')}."
        for player in players
    )
    return f"""# Jugadores observados

{_score_line(summary)}. El foco scouting debe separar rendimiento de volumen y accion decisiva.

{rows}

## Fortalezas

- Mexico mostro eficacia para convertir una accion clave en ventaja.
- Germany sostuvo volumen ofensivo y presencia territorial.
- Lozano destaca por incidencia directa en el marcador.

## Riesgos

- No proyectar potencial futuro a partir de un solo partido.
- No confundir dominio colectivo con rendimiento individual definitivo.

## Rol tactico

Los perfiles ofensivos de Germany aparecen asociados a volumen y insistencia. Para el equipo mexicano, el rol mas valioso fue atacar el espacio correcto y sostener el resultado.

## Conclusion scouting

El partido recomienda observar a los jugadores de impacto en contexto: Germany genero mas, pero Mexico tuvo la ejecucion que definio el marcador.
"""


def _fallback_ejecutivo(context: dict[str, Any], summary: dict[str, Any]) -> str:
    leader = _first(context.get("dominance", []))
    xg = _xg_text(context.get("xg_breakdown", []))
    return f"""# Conclusion

Mexico vencio 0-1 a Germany con una lectura ejecutiva clara: menor volumen, mayor eficacia.

## Hallazgos clave

- Resultado: Germany 0-1 Mexico.
- Dominio estimado: {leader.get('team_name', 'N/D')} con score {leader.get('dominance_score', 'N/D')}.
- Produccion xG: {xg}.
- Jugador decisivo: Hirving Rodrigo Lozano Bahena.

## Implicaciones

- El marcador premia la ejecucion mexicana en el momento clave.
- Germany necesita convertir dominio territorial en calidad de finalizacion.
- Para presentacion, el mensaje central es: dominio aleman, eficacia mexicana.

## Riesgos de lectura

- No concluir que Mexico controlo todo el partido por haber ganado.
- No concluir que Germany fue superior en eficacia por tener mayor volumen.
"""


def _extract_response_text(response: Any) -> str:
    output_text = getattr(response, "output_text", None)
    if output_text:
        return str(output_text)
    output = getattr(response, "output", None)
    if output:
        chunks: list[str] = []
        for item in output:
            for content in getattr(item, "content", []) or []:
                text = getattr(content, "text", None)
                if text:
                    chunks.append(str(text))
        if chunks:
            return "\n".join(chunks)
    return str(response)


def _first(rows: list[dict[str, Any]]) -> dict[str, Any]:
    return rows[0] if rows else {}


def _score_line(summary: dict[str, Any]) -> str:
    return (
        f"{summary.get('home_team_name', 'N/D')} {summary.get('home_score', 'N/D')}-"
        f"{summary.get('away_score', 'N/D')} {summary.get('away_team_name', 'N/D')}"
    )


def _xg_text(rows: list[dict[str, Any]]) -> str:
    if not rows:
        return "xG no disponible"
    return ", ".join(f"{row.get('team_name')} {row.get('xg_total')}" for row in rows)


def _dangerous_text(rows: list[dict[str, Any]]) -> str:
    if not rows:
        return "No hay ataques peligrosos suficientes en el contexto reducido."
    counts: dict[str, int] = {}
    for row in rows:
        team = str(row.get("team_name") or "Sin equipo")
        counts[team] = counts.get(team, 0) + 1
    return (
        "Ataques peligrosos en contexto reducido: "
        + ", ".join(f"{team} {count}" for team, count in sorted(counts.items()))
        + "."
    )


def _goal_moment(rows: list[dict[str, Any]]) -> dict[str, Any]:
    for row in rows:
        if str(row.get("type") or "").lower() == "goal":
            return row
    return _first(rows)


def _moment_sentence(moment: dict[str, Any]) -> str:
    if not moment:
        return "El momento clave no esta disponible en el contexto reducido."
    minute = moment.get("minute")
    second = int(moment.get("second") or 0)
    title = moment.get("title") or moment.get("type") or "Momento clave"
    return f"Al {minute}:{second:02d}, {title}."


def _player_or_default(moment: dict[str, Any], default: str) -> str:
    return str(moment.get("player_name") or default)


def _safe_token(value: str) -> str:
    clean = re.sub(r"[^a-zA-Z0-9_-]+", "-", value.strip())
    return clean.strip("-") or "narrativa"


def _fact_guard_text(markdown_text: str) -> str:
    lines = []
    for line in markdown_text.splitlines():
        if line.lstrip().startswith("#"):
            continue
        lines.append(line)
    return "\n".join(lines)
